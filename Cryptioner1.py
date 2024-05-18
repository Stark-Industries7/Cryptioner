# import string
# import random

# class Cryptioner:
#     def __init__(self,sttr,enkey):
#         self.sttr = sttr
#         self.enkey = enkey
#     def Encrypt(self):
#         return self.enkey.join(self.sttr)
#     def Decrypt(self):
#         return self.sttr.split(self.enkey)

# def random_enkey(length=5):
#     enkey = list(string.ascii_letters)
#     return ''.join(random.sample(enkey,length))


# def main():
#     print("------------------------Welcome To Cryptioner------------------------\n")
#     mainString = input("enter your string : ")
#     print()
    
#     encryptStr = random_enkey(5)
#     print("Your Random Generated Enkey is : ", encryptStr,"\n")

#     crypt = Cryptioner(mainString,encryptStr)
#     enStr = crypt.Encrypt()
#     deLst = crypt.Decrypt()
#     deStr = "".join(deLst)

#     print("The encrypted string is : ", enStr,"\n")
#     print("String after decryption is : ", deStr,"\n")
#     print("Thank you for using Cryptioner!😉 \n")
    
# if __name__ == "__main__":
#     main()


if __name__ == "__main__":
    import string
    import random
    import os
    from docx import Document

    class Cryptioner:
        def __init__(self, sttr, enkey):
            self.sttr = sttr
            self.enkey = enkey

        def Encrypt(self):
            return self.enkey.join(self.sttr)

        def Decrypt(self):
            return self.sttr.split(self.enkey)

    def generate_random_enkey(length=5):
        enkey_chars = list(string.ascii_letters)
        return ''.join(random.sample(enkey_chars, length))

    def read_text_file(file_path):
        with open(file_path, 'r') as file:
            return file.read()

    def write_text_file(file_path, content):
        with open(file_path, 'w') as file:
            file.write(content)

    def read_word_file(file_path):
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def write_word_file(file_path, content):
        doc = Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_path)

    def main():
        print("------------------------Welcome To Cryptioner------------------------\n")
    
        file_path = input("Enter the path of your file (text or word): ").strip()
        file_type = os.path.splitext(file_path)[1]
    
        random_enkey = generate_random_enkey(5)
        print(f"Generated encryption key: {random_enkey}\n")
    
        if file_type == '.txt':
            content = read_text_file(file_path)
        elif file_type == '.docx':
            content = read_word_file(file_path)
        else:
            print("Unsupported file type.")
            return
    
        crypt = Cryptioner(content, random_enkey)
        enStr = crypt.Encrypt()
    
        encrypted_file_path = file_path.replace(file_type, f"_encrypted{file_type}")
        if file_type == '.txt':
            write_text_file(encrypted_file_path, enStr)
        elif file_type == '.docx':
            write_word_file(encrypted_file_path, enStr)
    
        print(f"The encrypted file has been saved as: {encrypted_file_path}\n")
    
        # For demonstration, decrypt the file content
        deLst = crypt.Decrypt()
        deStr = "".join(deLst)
        decrypted_file_path = file_path.replace(file_type, f"_decrypted{file_type}")
    
        if file_type == '.txt':
            write_text_file(decrypted_file_path, deStr)
        elif file_type == '.docx':
            write_word_file(decrypted_file_path, deStr)
    
        print(f"The decrypted file has been saved as: {decrypted_file_path}\n")
        print("Thank you for using Cryptioner!😉 \n")

    if __name__ == "__main__":
        main()
